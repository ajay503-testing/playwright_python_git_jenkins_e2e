from docx.shared import Inches


def word_doc_utilty(document,brand,item,dates_pip,pip_screenshot,cart_screenshot,checkout_dates,checkout_screenshot,cart_dates,video_path):
    document.add_heading(f"{brand} - {item}", level=2)
    document.add_paragraph(f"PIP Dates: {dates_pip}")
    document.add_picture(pip_screenshot, width=Inches(3))
    document.add_paragraph(f"Cart Dates: {cart_dates}")
    document.add_picture(cart_screenshot, width=Inches(3))
    document.add_paragraph(f"Checkout Dates: {checkout_dates}")
    document.add_picture(checkout_screenshot, width=Inches(3))
    document.add_paragraph(f"Video Path: {video_path}")
    document.add_page_break()