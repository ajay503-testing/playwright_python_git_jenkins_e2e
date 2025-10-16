import os

from docx import Document
from docx.shared import Inches
from playwright.sync_api import Playwright
from playwright.sync_api import Page
from configparser import ConfigParser
from docx.shared import Pt
import utils.utils_actions_test
import time

test_data_list=[
    {"Brand":"PB","item":"6473869"},
    {"Brand":"WS","item":"4057363"},
    #{"Brand":"WS","item":"100182"}
]

os.makedirs("screenshots", exist_ok=True)
os.makedirs("videos", exist_ok=True)

def test_capture_dates(playwright:Playwright):
    document = Document()
    style = document.styles["Normal"]
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    document.add_heading("AODD Dates Capture Report", level=1)
    for index, test_data in enumerate(test_data_list):
        brand = test_data["Brand"]
        item = test_data["item"]

        print(f"===== Running dataset {index + 1}: {brand} - {item} =====")

        output_dir = f"screenshots/{brand}_{item}"
        os.makedirs(output_dir, exist_ok=True)

        browser = None

        try:
            browser = playwright.chromium.launch(headless=False)
            context = browser.new_context()
            page = context.new_page()

            utils.utils_actions_test.launchURL(page, test_data["Brand"])
            time.sleep(15)

            utils.utils_actions_test.closeDialog(page,test_data["Brand"])
            print("Dialog is closed")

            utils.utils_actions_test.enterProduct(page, test_data["item"])
            print("Product is entered")

            utils.utils_actions_test.enterPostalCode(page)
            print("Postal code is entered")

            time.sleep(10)
            dates_pip = utils.utils_actions_test.getAODDDatesOnPipPage(page)
            print(dates_pip)
            pip_screenshot = f"{output_dir}/pip_page.png"
            page.screenshot(path=pip_screenshot, full_page=True)
            print("Captured PIP Page Screenshot")

            utils.utils_actions_test.addItemToCart(page)
            cart_dates = utils.utils_actions_test.getAODDDatesonCartPage(page)
            print("AODD Dates on Cart Page:", cart_dates)
            cart_screenshot = f"{output_dir}/cart_page.png"
            page.screenshot(path=cart_screenshot,full_page=True)
            print("Captured Cart Page Screenshot")

            checkout_dates = utils.utils_actions_test.addProductToCheckOut(page,test_data["Brand"])
            print("Checkout Dates:", checkout_dates)
            checkout_screenshot = f"{output_dir}/checkout_page.png"
            page.screenshot(path=checkout_screenshot,full_page=True)
            print("Captured Checkout Page Screenshot")

            time.sleep(10)
            # Wait for recording to finalize
            time.sleep(5)
            video_path = context.pages[0].video.path() if context.pages[0].video else "No video"

            # Add results to Word doc
            document.add_heading(f"{brand} - {item}", level=2)
            document.add_paragraph(f"PIP Dates: {dates_pip}")
            document.add_picture(pip_screenshot,  width=Inches(3))
            document.add_paragraph(f"Cart Dates: {cart_dates}")
            document.add_picture(cart_screenshot,  width=Inches(3))
            document.add_paragraph(f"Checkout Dates: {checkout_dates}")
            document.add_picture(checkout_screenshot,  width=Inches(3))
            document.add_paragraph(f"Video Path: {video_path}")
            document.add_page_break()

        except Exception as e:
            print(f"Error in iteration {index}: {e}")

        finally:
            # Always close context and browser
            try:
                context.close()
            except Exception:
                pass
            if browser:
                browser.close()
            print(f"Browser closed for iteration {index}")
    report_path = "AODD_Date_Report.docx"
    document.save(report_path)
    print(f"\n✅ Report saved successfully: {report_path}")








